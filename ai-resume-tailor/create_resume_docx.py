# create_resume_docx.py

import json
import logging
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)


def add_horizontal_line(paragraph):
    """
    Applies a single bottom border to the given paragraph, creating a horizontal line.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')          # thickness of the line
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')  # black
    pBdr.append(bottom)
    pPr.append(pBdr)
    logger.debug("Added horizontal line to paragraph.")


def make_two_column_paragraph(doc, left_text, right_text, bold_left=False, italic_left=False):
    """
    Creates one paragraph with a right-aligned tab stop.
    The left_text (optionally bold/italic) appears at the left;
    the right_text is pushed to the right margin.
    """
    para = doc.add_paragraph()
    para_format = para.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(0)
    para_format.line_spacing = 1

    tab_stops = para_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.5), alignment=WD_TAB_ALIGNMENT.RIGHT)

    run_left = para.add_run(left_text)
    if bold_left:
        run_left.bold = True
    if italic_left:
        run_left.italic = True

    para.add_run(f'\t{right_text}')
    logger.debug(f"Created two-column paragraph: '{left_text}' - '{right_text}'.")
    return para


def add_bulleted_item(doc, text):
    """
    Adds a single bullet item with a fixed indent of 0.25".
    """
    para = doc.add_paragraph(style='List Bullet')
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.25)
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(0)
    para_format.line_spacing = 1

    run = para.add_run(text)
    logger.debug(f"Added bulleted item: '{text[:60]}...'")
    return para


def generate_docx_from_json(resume_data: dict, output_path: str = "Praneeth_Ravuri_Tightened.docx"):
    """
    Given a resume_data dictionary (parsed from JSON), build and save a .docx file
    named output_path. resume_data must have keys: "header", "work_experience",
    "education", "skills", "projects".
    """
    logger.info(f"Generating DOCX at '{output_path}'.")
    doc = Document()

    # -----------------------------
    # 1) Adjust page margins
    # -----------------------------
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    logger.debug("Page margins set: left/right=0.5in, top/bottom=0.25in.")

    # -----------------------------
    # 2) Global default font & paragraph spacing
    # -----------------------------
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1
    logger.debug("Global default font set to Calibri 10pt, spacing adjusted.")

    # -----------------------------
    # 3) HEADER (Name + Contact)
    # -----------------------------
    header = resume_data.get("header", {})
    name_para = doc.add_paragraph()
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(0)
    run = name_para.add_run(header.get("name", ""))
    run.bold = True
    run.font.size = Pt(16)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(4)
    contact_info = " | ".join(header.get("contact", []))
    run = contact_para.add_run(contact_info)
    run.font.size = Pt(10)
    logger.debug("Header (name + contact) added.")

    # -----------------------------
    # 4) WORK EXPERIENCE
    # -----------------------------
    we_title = doc.add_paragraph()
    we_title.alignment = WD_TAB_ALIGNMENT.LEFT
    we_title.paragraph_format.space_before = Pt(2)
    we_title.paragraph_format.space_after = Pt(0)
    run = we_title.add_run("WORK EXPERIENCE")
    run.bold = True
    run.font.size = Pt(12)
    add_horizontal_line(we_title)

    for exp in resume_data.get("work_experience", []):
        company = exp.get("company", "")
        location = exp.get("location", "")
        position = exp.get("position", "")
        start_date = exp.get("start_date", "")
        end_date = exp.get("end_date", "")

        p1 = make_two_column_paragraph(
            doc,
            left_text=company,
            right_text=location,
            bold_left=True
        )
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(0)

        p2 = make_two_column_paragraph(
            doc,
            left_text=position,
            right_text=f"{start_date} – {end_date}",
            bold_left=True
        )
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)

        for bullet in exp.get("bullets", []):
            b = add_bulleted_item(doc, bullet)
            b.paragraph_format.space_before = Pt(0)
            b.paragraph_format.space_after = Pt(0)

    # -----------------------------
    # 5) EDUCATION
    # -----------------------------
    edu_title = doc.add_paragraph()
    edu_title.alignment = WD_TAB_ALIGNMENT.LEFT
    edu_title.paragraph_format.space_before = Pt(4)
    edu_title.paragraph_format.space_after = Pt(0)
    run = edu_title.add_run("EDUCATION")
    run.bold = True
    run.font.size = Pt(12)
    add_horizontal_line(edu_title)

    for edu in resume_data.get("education", []):
        institution = edu.get("institution", "")
        location = edu.get("location", "")
        degree = edu.get("degree", "")
        start_date = edu.get("start_date", "")
        end_date = edu.get("end_date", "")

        p3 = make_two_column_paragraph(
            doc,
            left_text=institution,
            right_text=location,
            bold_left=True
        )
        p3.paragraph_format.space_before = Pt(4)
        p3.paragraph_format.space_after = Pt(0)

        p4 = make_two_column_paragraph(
            doc,
            left_text=degree,
            right_text=f"{start_date} – {end_date}",
        )
        p4.paragraph_format.space_before = Pt(0)
        p4.paragraph_format.space_after = Pt(4)

    # -----------------------------
    # 6) SKILLS
    # -----------------------------
    skills_title = doc.add_paragraph()
    skills_title.alignment = WD_TAB_ALIGNMENT.LEFT
    skills_title.paragraph_format.space_before = Pt(4)
    skills_title.paragraph_format.space_after = Pt(4)
    run = skills_title.add_run("SKILLS")
    run.bold = True
    run.font.size = Pt(12)
    add_horizontal_line(skills_title)

    for skill_group in resume_data.get("skills", []):
        name = skill_group.get("name", "")
        items = skill_group.get("items", [])
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1
        para.add_run(f"{name}: " + ", ".join(items))

    # -----------------------------
    # 7) PROJECTS
    # -----------------------------
    proj_title = doc.add_paragraph()
    proj_title.alignment = WD_TAB_ALIGNMENT.LEFT
    proj_title.paragraph_format.space_before = Pt(4)
    proj_title.paragraph_format.space_after = Pt(0)
    run = proj_title.add_run("PROJECTS")
    run.bold = True
    run.font.size = Pt(12)
    add_horizontal_line(proj_title)

    for proj in resume_data.get("projects", []):
        project_name = proj.get("name", "")
        bullets = proj.get("bullets", [])

        p5 = doc.add_paragraph()
        p5.paragraph_format.space_before = Pt(4)
        p5.paragraph_format.space_after = Pt(4)
        p5.paragraph_format.line_spacing = 1
        run = p5.add_run(project_name)
        run.bold = True

        for bullet in bullets:
            pb = doc.add_paragraph(style='List Bullet')
            pb.paragraph_format.left_indent = Inches(0.25)
            pb.paragraph_format.space_before = Pt(0)
            pb.paragraph_format.space_after = Pt(0)
            pb.paragraph_format.line_spacing = 1
            pb.add_run(bullet)

    # -----------------------------
    # 8) Save the Document
    # -----------------------------
    doc.save(output_path)
    logger.info(f"✅ Document generated: {output_path}")
